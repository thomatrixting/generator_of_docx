from docx import Document
from importlib_metadata import sys
import pandas as pd


def filter_df (df):

    df.columns = [''] * len(df.columns)
    colums_list = df.iloc[[2]].values.tolist()[0]
    df.columns  = [value.replace(' ','').lower() for value in colums_list] #set colums

    df = df[(df.iloc[:, 0] == df.iloc[:, 0])] #eliminate the nan
    df = df.drop([2]) #eliminate row with colums names
    return df 

def main(relative_path_excel,path_to_save_words):
    try:
        df = pd.read_excel(relative_path_excel)
        df = filter_df(df)
    except Exception as error:
        error_read = input('error al leer excel quieres ver el traceback(y): ')
        if error_read.startswith('y'):
            print(error)
        sys.exit()
    df = df.rename(columns={'nonbre':'name','cedula':'id'})

    for index,row in df.iterrows():
        try:
            doc = Document('templeates/certificate_templeate.docx')
        except Exception as error:
            error_read = input('error al leer el templeate quieres ver el traceback(y): ')
            if error_read.startswith('y'):
                print(error)
            sys.exit()
            
        

        text_to_change = doc.paragraphs[16].text 
        text_to_change = text_to_change.format(row['name'],row['id'])
        doc.paragraphs[16].text  = text_to_change 
        
        try:
            name = row['name'].replace(' ','_')
            doc.save(f'{path_to_save_words}/{name}_certificacion.docx')
        except Exception as error:
            error_read = input('error al guardar los archivos quieres ver el traceback(y): ')
            if error_read.startswith('y'):
                print(error)
            sys.exit()

def generate_imput():
    file_reference_path = input('entra la ubicacion del exel: ')
    directory_to_safe = input('entra el directorio donde van a quedar los archivos: ')
    main(file_reference_path,directory_to_safe)


def run():
    try:
        generate_imput()
    except Exception as error:
        error_read = input('error desconocido quieres ver el traceback(y): ')
        if error_read.startswith('y'):
            print(error)
        sys.exit()

if __name__ == '__main__':
    run()