from docx import Document

doc = Document('templeates/certificate_templeate.docx')

# counter = 0
# for paragraph in doc.paragraphs: #is the 16 paragraph
#     print(paragraph.text)
#     print(counter)
#     counter += 1

text = doc.paragraphs[16].text
text = text.format('thomas','1031650')
doc.paragraphs[16].text = text

# print(text)
print(doc.paragraphs[16].text)

doc.save('files_words/demo.docx')


